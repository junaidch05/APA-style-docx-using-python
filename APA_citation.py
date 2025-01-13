import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Load data from the JSON file
with open('cite.json', 'r', encoding='utf-8') as file:
    details = json.load(file)


def citation(doc_name,details):


    # Create a new Document
    doc = Document()

    # Set the default font style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Function to set double spacing
    def set_double_spacing(paragraph):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = Pt(24)  # Double space



    # Add page number in the upper right corner
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add a field for page numbers
    page_number = OxmlElement('w:fldSimple')
    page_number.set(qn('w:instr'), 'PAGE')
    paragraph._element.append(page_number)
    paragraph.add_run(' ')

    # Title of the Paper (centered, bold, 1/3 down the page


    # Add "References" title (centered, bold)
    references_title = doc.add_paragraph()
    references_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    references_run = references_title.add_run('Citation')
    references_run.bold = True
    references_run.font.size = Pt(12)
    set_double_spacing(references_title)

    # Function to set hanging indentation
    def set_hanging_indent(paragraph, indent_in_inches=0.5):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Inches(indent_in_inches)
        paragraph_format.first_line_indent = Inches(-indent_in_inches)

    # Function to format citation text in APA style
    def apply_italic(run):
        run.italic = True

    def format_citation(citation, doc):
        citation_paragraph = doc.add_paragraph()
        citation_type = citation['type']
        
        if citation_type == 'book':
            authors = citation['author']
            for author in authors:
                citation_paragraph.add_run(f"{author['last_name']}, {author['first_name'][0]}. ")
            citation_paragraph.add_run(f"({citation['year']}). ")
            title_run = citation_paragraph.add_run(citation['title'])
            apply_italic(title_run)
            citation_paragraph.add_run(f". {citation['publisher']}.")
        elif citation_type == 'journal_article':
            authors = citation['author']
            for author in authors:
                citation_paragraph.add_run(f"{author['last_name']}, {author['first_name'][0]}. ")
            citation_paragraph.add_run(f"({citation['year']}). ")
            title_run = citation_paragraph.add_run(citation['title'])
            apply_italic(title_run)
            citation_paragraph.add_run(f". {citation['journal']}, {citation['volume']}({citation['issue']}), {citation['pages']}. {citation.get('doi', '')}")
        elif citation_type == 'website':
            if citation['author']:
                authors = citation['author']
                for author in authors:
                    citation_paragraph.add_run(f"{author['last_name']}, {author['first_name'][0]}. ")
            citation_paragraph.add_run(f"({citation['date']}). ")
            title_run = citation_paragraph.add_run(citation['title'])
            apply_italic(title_run)
            citation_paragraph.add_run(f". {citation['website_name']}. {citation['url']}") 
        elif citation_type == 'newspaper_article':
            citation_paragraph.add_run(f"{citation['author'][0]['last_name']}, {citation['author'][0]['first_name'][0]}. ")
            citation_paragraph.add_run(f"({citation['date']}). {citation['title']}. {citation['newspaper']}. {citation['url']}")
        elif citation_type == 'book_chapter':
            citation_paragraph.add_run(f"{citation['author'][0]['last_name']}, {citation['author'][0]['first_name'][0]}. ")
            citation_paragraph.add_run(f"({citation['year']}). {citation['chapter_title']}. In ")
            citation_paragraph.add_run(f"{citation['editor'][0]['last_name']}, {citation['editor'][0]['first_name'][0]} (Ed.), ")
            book_title_run = citation_paragraph.add_run(citation['book_title'])
            apply_italic(book_title_run)
            citation_paragraph.add_run(f" (pp. {citation['pages']}). {citation['publisher']}.")

        return citation_paragraph

    # Add citations to the document
    for citation in details['citations']:
        citation_paragraph = format_citation(citation, doc)
        set_double_spacing(citation_paragraph)
        set_hanging_indent(citation_paragraph)

    # Save the document
    doc.save(doc_name)


if __name__ == "__main__":
    citation("doc_name",details)
