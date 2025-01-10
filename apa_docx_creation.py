import json
import requests
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
import os

with open('docx1.json', 'r', encoding='utf-8') as file:
    details = json.load(file)

def document_execution(details):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_height = Inches(11)  # Set height to 11 inches
        section.page_width = Inches(8.5)   # Set width to 8.5 inches


    # Set the default font style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Function to set double spacing
    def set_double_spacing(paragraph):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = Pt(24)  # Double space

    # Function to set first-line indentation
    def set_first_line_indent(paragraph, indent_in_inches=0.5):
        paragraph.paragraph_format.first_line_indent = Inches(indent_in_inches)

    # Function to set block quote formatting
    def set_block_quote_formatting(paragraph):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Inches(0.5)  # Indent block quote
        set_double_spacing(paragraph)


    def set_hanging_indent(paragraph, indent_in_inches=0.5):
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Inches(indent_in_inches)
            paragraph_format.first_line_indent = Inches(-indent_in_inches)

    # Function to add lists
    def add_list(doc, list_content):
        if list_content.get('list_type') == 'unordered':
            for item in list_content['items']:
                p = doc.add_paragraph(item, 'List Bullet')
                p.paragraph_format.left_indent = Inches(0.5)
                set_double_spacing(p)
        elif list_content.get('list_type') == 'ordered':
            for item in list_content['items']:
                p = doc.add_paragraph(item, 'List Number')
                p.paragraph_format.left_indent = Inches(0.5)
                set_double_spacing(p)



    def  add_figures_to_document(doc, figures_data):
        folder_path = os.path.dirname(os.path.abspath(__file__))
        figure_number = figures_data['figure_number']
        figure_title = figures_data['figure_title']
    
        # Add a paragraph for the table number and make it bold
        figure_number_paragraph = doc.add_paragraph()
        run = figure_number_paragraph.add_run(figure_number)
        run.bold = True
        set_double_spacing(figure_number_paragraph)  # Apply double spacing

        # Add a paragraph for the table title (italicized)
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(figure_title)
        title_run.italic = True
        set_double_spacing(title_paragraph)  # Apply double spacing
        image_path = figures_data['figure_url']
        
        if image_path.startswith('http://') or image_path.startswith('https://'):
            # Fetch image from URL
            try:
                response = requests.get(image_path)
                response.raise_for_status()  # Raise an error for bad responses
                image = BytesIO(response.content)
                doc.add_picture(image, width=Inches(6), height=Inches(4))
            except requests.exceptions.RequestException as e:
                print(f"Error fetching image: {e}")
        else:
            # Use local image path
            local_image_path = os.path.join(folder_path, image_path)
            try:
                doc.add_picture(local_image_path, width=Inches(6), height=Inches(4))
            except Exception as e:
                print(f"Error adding local image: {e}")

        if "notes" in figures_data:
            notes_paragraph = doc.add_paragraph()
            # Make the first word italic
            first_word = figures_data["notes"].split()[0]  # Get the first word
            first_word_run = notes_paragraph.add_run(first_word)
            first_word_run.italic = True

            # Add the rest of the notes after the first word
            rest_of_notes = ' '.join(figures_data["notes"].split()[1:])  # Get the rest of the notes
            notes_paragraph.add_run(' ' + rest_of_notes)  # Add a space before the rest

            notes_paragraph.style = doc.styles['Normal']
            set_double_spacing(notes_paragraph)
        doc.add_paragraph() 



    def add_table_to_document(doc, table_data):
        # Add table number and title
        table_number = table_data['table_number']
        table_title = table_data['table_title']
        
        # Add a paragraph for the table number and make it bold
        table_number_paragraph = doc.add_paragraph()
        run = table_number_paragraph.add_run(table_number)
        run.bold = True
        set_double_spacing(table_number_paragraph)  # Apply double spacing

        # Add a paragraph for the table title (italicized)
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(table_title)
        title_run.italic = True
        set_double_spacing(title_paragraph)  # Apply double spacing
    
        # Add the table with headers
        table_paragraph = doc.add_paragraph()
        table_paragraph.paragraph_format.space_before = Pt(0)  
        table = doc.add_table(rows=1, cols=len(table_data['headers']))

        
        # Set headers and align them
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(table_data['headers']):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
        
        # Add data rows
        for row_data in table_data['data']:
            row_cells = table.add_row().cells
            for i, item in enumerate(row_data):
                row_cells[i].text = str(item)
                if i == 0:  # Left align the first column data
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:  # Center align other columns
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Apply borders: Horizontal lines only, no vertical lines
        tbl = table._element
        tbl_borders = parse_xml(
            r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            r'<w:top w:val="single" w:sz="6"/>'   # Top border above the table
            r'<w:bottom w:val="single" w:sz="6"/>' # Bottom border below the table
            r'<w:insideH w:val="single" w:sz="6"/>' # Horizontal line between header and data
            r'<w:insideV w:val="nil"/>' # No vertical borders
            r'</w:tblBorders>'
        )
        tbl.tblPr.append(tbl_borders)
        table_paragraph = doc.add_paragraph()
    # Space above the table
        table_paragraph.paragraph_format.space_after = Pt(0) 

        # No extra blank space; table naturally takes space like a paragraph

        # Add notes (if any) below the table
        if "notes" in table_data:
            notes_paragraph = doc.add_paragraph()
            # Make the first word italic
            first_word = table_data["notes"].split()[0]  # Get the first word
            first_word_run = notes_paragraph.add_run(first_word)
            first_word_run.italic = True

            # Add the rest of the notes after the first word
            rest_of_notes = ' '.join(table_data["notes"].split()[1:])  # Get the rest of the notes
            notes_paragraph.add_run(' ' + rest_of_notes)  # Add a space before the rest

            notes_paragraph.style = doc.styles['Normal']
            set_double_spacing(notes_paragraph)
        doc.add_paragraph()      # Apply double spacing for the notes

    def apply_italic(run):
            run.italic = True


    # Add page number in the upper right corner
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add a field for page numbers
    page_number = OxmlElement('w:fldSimple')
    page_number.set(qn('w:instr'), 'PAGE')
    paragraph._element.append(page_number)
    paragraph.add_run(' ')

    # Title of the Paper (centered, bold, 1/3 down the page)
    for _ in range(6):  # Add space to move title down
        doc.add_paragraph()

    if 'title' in details:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(details['title'])
        run.bold = True
        run.font.size = Pt(12)
        set_double_spacing(title)
        doc.add_paragraph()

    # Add space after the title
    # Author's Name
    if 'author_name' in details:
        author = doc.add_paragraph()
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author.add_run(details['author_name'])
        set_double_spacing(author)

    # Institutional Affiliation
    if 'affiliation' in details:
        affiliation = doc.add_paragraph()
        affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
        affiliation_run = affiliation.add_run(details['affiliation'])
        set_double_spacing(affiliation)

    # Course Number
    # Institutional Affiliation
    if 'course_number' in details:
        course = doc.add_paragraph()
        course.alignment = WD_ALIGN_PARAGRAPH.CENTER
        course_run = course.add_run(details['course_number'])
        set_double_spacing(course)

    # Instructor's Name
    if 'instructor_name' in details:
        instructor = doc.add_paragraph()
        instructor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        instructor_run = instructor.add_run(details['instructor_name'])
        set_double_spacing(instructor)

    # Date
    if 'date'in details:
        date = doc.add_paragraph()
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date.add_run(details['date'])
        set_double_spacing(date)

    # Add a page break to move to the abstract page
    doc.add_page_break()

    # Abstract Title (centered, bold)

    # Add abstract text
    if 'abstract'in details:
        abstract_title = doc.add_paragraph()
        abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abstract_run = abstract_title.add_run('Abstract')
        abstract_run.bold = True
        abstract_run.font.size = Pt(12)
        for abstract_paragraph in details['abstract']:
            paragraph = doc.add_paragraph(abstract_paragraph)
            set_double_spacing(paragraph)

    # Keywords
    if 'keywords' in details:
        keywords_paragraph = doc.add_paragraph()
        set_first_line_indent(keywords_paragraph)
        keywords_str = ', '.join(details['keywords'])
        keywords_paragraph.add_run(f"Keywords: {keywords_str}")
        set_double_spacing(keywords_paragraph)

    # Add a page break to move to the main sections
    doc.add_page_break()

    # Function to add headings with optional content
    def add_heading(doc, text, level, content=None):
        styles = {
            1: {'bold': True, 'size': Pt(12), 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
            2: {'bold': True, 'size': Pt(12), 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
            3: {'bold': True, 'italic': True, 'size': Pt(12), 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
            4: {'bold': True, 'size': Pt(12), 'alignment': WD_ALIGN_PARAGRAPH.LEFT, 'indent': Inches(0.5)},
            5: {'bold': True, 'italic': True, 'size': Pt(12), 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
        }

        style_name = f'Heading{level}'

        if style_name not in doc.styles:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.bold = styles[level].get('bold', False)
            style.font.italic = styles[level].get('italic', False)
            style.font.size = styles[level]['size']
            style.paragraph_format.alignment = styles[level]['alignment']

            if 'indent' in styles[level]:
                style.paragraph_format.first_line_indent = styles[level]['indent']

        if level == 4 and content:
            heading_paragraph = doc.add_paragraph()
            run = heading_paragraph.add_run(text + ". ")
            run.bold = True

            first_content = content[0]
            content_run = heading_paragraph.add_run(first_content['text'])

            set_first_line_indent(heading_paragraph)
            set_double_spacing(heading_paragraph)

            for paragraph_content in content[1:]:
                if isinstance(paragraph_content, dict) and 'block' in paragraph_content:
                    block_content = paragraph_content['block']
                    if 'quote' in block_content and block_content['quote']:
                        block_paragraph = doc.add_paragraph(block_content['quote'])
                        set_block_quote_formatting(block_paragraph)
                    if 'following_text' in block_content and block_content['following_text']:
                        following_paragraph = doc.add_paragraph(block_content['following_text'])
                        set_double_spacing(following_paragraph)  
                
                elif isinstance(paragraph_content, dict) and 'list' in paragraph_content:
                    list_content = paragraph_content['list']
                    add_list(doc, list_content)
                
                elif isinstance(paragraph_content, dict) and 'table_data' in paragraph_content:
                    table_content = paragraph_content['table_data']
                    add_table_to_document(doc, table_content)
                
                elif isinstance(paragraph_content, dict) and 'figures' in paragraph_content:
                            figures_content = paragraph_content['figures']
                            add_figures_to_document(doc, figures_content)
                else:
                    content_paragraph = doc.add_paragraph(paragraph_content['text'])
                    set_double_spacing(content_paragraph)
                    set_first_line_indent(content_paragraph)
        else:
            paragraph = doc.add_paragraph(text, style=style_name)
            set_double_spacing(paragraph)

    # Add Sections from JSON
    if 'sections' in details:
        for section in details['sections']:
            heading = section['heading']
            level = section['level']
            content = section['content']

            if level == 4:
                add_heading(doc, heading, level, content=content)
            else:
                add_heading(doc, heading, level)
                if content:
                    for paragraph_content in content:
                        if isinstance(paragraph_content, dict) and 'block' in paragraph_content:
                            block_content = paragraph_content['block']
                            if 'quote' in block_content:
                                block_paragraph = doc.add_paragraph(block_content['quote'])
                                set_block_quote_formatting(block_paragraph)
                            if 'following_text' in block_content:
                                following_paragraph = doc.add_paragraph(block_content['following_text'])
                                set_double_spacing(following_paragraph)
                        
                        elif isinstance(paragraph_content, dict) and 'list' in paragraph_content:
                            list_content = paragraph_content['list']
                            add_list(doc, list_content)

                        elif isinstance(paragraph_content, dict) and 'table_data' in paragraph_content:
                            table_content = paragraph_content['table_data']
                            if table_content:  # Check if table_content is not empty
                                add_table_to_document(doc, table_content)

                        elif isinstance(paragraph_content, dict) and 'figures' in paragraph_content:
                            figures_content = paragraph_content['figures']
                            if figures_content:
                                add_figures_to_document(doc, figures_content)
                        
                        else:
                            content_paragraph = doc.add_paragraph(paragraph_content['text'])
                            set_double_spacing(content_paragraph)
                            set_first_line_indent(content_paragraph)

    # Add a page break before the reference section


    # Function to set hanging indentation


    def format_citation(citation, doc):
        citation_paragraph = doc.add_paragraph()
        citation_type = citation['type']
        
        if citation_type == 'book':
            authors = citation['author']
            for author in authors:
                citation_paragraph.add_run(f"{author['last_name']}, {author['first_name'][0]}. ")
            citation_paragraph.add_run(f"({citation['year']}). ")
            title_run = citation_paragraph.add_run(citation['title'])
            apply_italic(title_run)
            citation_paragraph.add_run(f". {citation['publisher']}.")
        elif citation_type == 'journal_article':
            authors = citation['author']
            for author in authors:
                citation_paragraph.add_run(f"{author['last_name']}, {author['first_name'][0]}. ")
            citation_paragraph.add_run(f"({citation['year']}). ")
            title_run = citation_paragraph.add_run(citation['title'])
            apply_italic(title_run)
            citation_paragraph.add_run(f". {citation['journal']}, {citation['volume']}({citation['issue']}), {citation['pages']}. {citation.get('doi', '')}")
        elif citation_type == 'website':
            if citation['author']:
                authors = citation['author']
                for author in authors:
                    citation_paragraph.add_run(f"{author['last_name']}, {author['first_name'][0]}. ")
            citation_paragraph.add_run(f"({citation['date']}). ")
            title_run = citation_paragraph.add_run(citation['title'])
            apply_italic(title_run)
            citation_paragraph.add_run(f". {citation['website_name']}. {citation['url']}") 
        elif citation_type == 'newspaper_article':
            citation_paragraph.add_run(f"{citation['author'][0]['last_name']}, {citation['author'][0]['first_name'][0]}. ")
            citation_paragraph.add_run(f"({citation['date']}). {citation['title']}. {citation['newspaper']}. {citation['url']}")
        elif citation_type == 'book_chapter':
            citation_paragraph.add_run(f"{citation['author'][0]['last_name']}, {citation['author'][0]['first_name'][0]}. ")
            citation_paragraph.add_run(f"({citation['year']}). {citation['chapter_title']}. In ")
            citation_paragraph.add_run(f"{citation['editor'][0]['last_name']}, {citation['editor'][0]['first_name'][0]} (Ed.), ")
            book_title_run = citation_paragraph.add_run(citation['book_title'])
            apply_italic(book_title_run)
            citation_paragraph.add_run(f" (pp. {citation['pages']}). {citation['publisher']}.")

        return citation_paragraph

    # Add citations to the document
    if 'citations' in details:
        doc.add_page_break()
        references_title = doc.add_paragraph()
        references_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        references_run = references_title.add_run('References')
        references_run.bold = True
        references_run.font.size = Pt(12)
        set_double_spacing(references_title)

        for citation in details['citations']:
            citation_paragraph = format_citation(citation, doc)
            set_double_spacing(citation_paragraph)
            set_hanging_indent(citation_paragraph)

        
    if 'tables_figure' in details:
        for content in details['tables_figure']:
                if isinstance(content, dict) and 'table_data' in content:
                            doc.add_page_break()
                            table_content = content['table_data']
                            add_table_to_document(doc, table_content)

                elif isinstance(content, dict) and 'figures' in content:
                            doc.add_page_break()
                            figures_content = content['figures']
                            add_figures_to_document(doc, figures_content)


    # Save the document
    doc.save('APA_Style_using_python.docx')


if __name__ == "__main__":
    document_execution(details)